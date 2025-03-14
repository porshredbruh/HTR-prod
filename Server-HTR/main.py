import os
import random
import math
from io import BytesIO

import numpy as np
import torch
from torch import nn
from torch.nn import Conv2d, MaxPool2d, BatchNorm2d, LeakyReLU
from torchvision import transforms
from PIL import Image, ImageDraw, ImageFont
from torch.utils.data import Dataset
import torchvision
from torchvision.models.detection.faster_rcnn import FastRCNNPredictor
import torchvision.transforms as T
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import FileResponse
from pathlib import Path

# Папки для загрузки и сохранения файлов
UPLOAD_FOLDER = './test/'
OUTPUT_FOLDER = './res/'
WEIGHTS_PATH_TEXT = "./model_ocr.pt"   # весы для обычного текста
WEIGHTS_PATH_MATH = "./model_math.pt"  # весы для математической формулы
model_path_faster = './model_faster.pth'

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

DEVICE = 'cuda:0' if torch.cuda.is_available() else 'cpu'

# ---------------------------------------------------------------------------------
# Алфавиты:
# ---------------------------------------------------------------------------------
ALPHABET_TEXT = [
    'PAD', 'SOS', ' ', '!', '"', '%', '(', ')', ',', '-', '.', '/',
    '0', '1', '2', '3', '4', '5', '6', '7', '8', '9', ':', ';', '?',
    '[', ']', '«', '»', 'А', 'Б', 'В', 'Г', 'Д', 'Е', 'Ж', 'З', 'И',
    'Й', 'К', 'Л', 'М', 'Н', 'О', 'П', 'Р', 'С', 'Т', 'У', 'Ф', 'Х',
    'Ц', 'Ч', 'Ш', 'Щ', 'Э', 'Ю', 'Я', 'а', 'б', 'в', 'г', 'д', 'е',
    'ж', 'з', 'и', 'й', 'к', 'л', 'м', 'н', 'о', 'п', 'р', 'с', 'т',
    'у', 'ф', 'х', 'ц', 'ч', 'ш', 'щ', 'ъ', 'ы', 'ь', 'э', 'ю', 'я',
    'ё', 'EOS'
]

ALPHABET_MATH = [
    'PAD', 'SOS', 'EOS', '0', '1', '2', '3', '4', '5', '6', '7', '8', '9',
    'a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n',
    'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z',
    'A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M', 'N',
    'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z', '«', '»',
    'А', 'Б', 'В', 'Г', 'Д', 'Е', 'Ж', 'З', 'И', 'Й', 'К', 'Л', 'М', 'Н',
    'О', 'П', 'Р', 'С', 'Т', 'У', 'Ф', 'Х', 'Ц', 'Ч', 'Ш', 'Щ', 'Э', 'Ю',
    'Я', 'а', 'б', 'в', 'г', 'д', 'е', 'ж', 'з', 'и', 'й', 'к', 'л', 'м',
    'н', 'о', 'п', 'р', 'с', 'т', 'у', 'ф', 'х', 'ц', 'ч', 'ш', 'щ', 'ъ',
    'ы', 'ь', 'э', 'ю', 'я', 'ё', '+', '-', '=', '(', ')', '[', ']', '{',
    '}', '\\', '/', '|', '<', '>', '.', ',', ':', '!', '?', '&', '%', '$',
    '#', '@', '*', '~', '`', '^', '_'
]

# ---------------------------------------------------------------------------------
# Faster R-CNN для сегментации (Text / Math / Image)
# ---------------------------------------------------------------------------------
CLASS_NAMES = {
    1: "Text",    # Класс 1: "Text"
    2: "Math",    # Класс 2: "Math"
    3: "Image",   # Класс 3: "Image"
}
COLORS = {
    1: "blue",   # Text
    2: "red",    # Math
    3: "green",  # Image
}

num_classes = 4  # 3 класса + фон

model_faster = torchvision.models.detection.fasterrcnn_resnet50_fpn(pretrained=False)
in_features = model_faster.roi_heads.box_predictor.cls_score.in_features
model_faster.roi_heads.box_predictor = FastRCNNPredictor(in_features, num_classes)

def load_model_faster(model, path, device):
    if not os.path.exists(path):
        raise FileNotFoundError(f"Faster R-CNN не найден по пути: {path}")
    checkpoint = torch.load(path, map_location=device, weights_only=True)
    model.load_state_dict(checkpoint)
    model.to(device)
    model.eval()
    print("Faster R-CNN загружен и готов к использованию.")
    return model

try:
    model_faster = load_model_faster(model_faster, model_path_faster, DEVICE)
except Exception as e:
    print("Ошибка загрузки Faster R-CNN:", e)
    exit(1)

# ---------------------------------------------------------------------------------
# Класс датасета с одним изображением
# ---------------------------------------------------------------------------------
class TestImageDataset(Dataset):
    def __init__(self, image_path, transforms=None):
        self.image_path = image_path
        self.transforms = transforms
        self.image_extensions = ['.jpg', '.jpeg', '.png', '.bmp', '.tif', '.tiff']
        
        if not os.path.isfile(self.image_path):
            raise FileNotFoundError(f"Файл {self.image_path} не найден.")
        ext = os.path.splitext(self.image_path)[1].lower()
        if ext not in self.image_extensions:
            raise ValueError(f"Неподдерживаемый формат файла: {ext}")
        
    def __len__(self):
        return 1
    
    def __getitem__(self, idx):
        if idx != 0:
            raise IndexError("В этом датасете только одно изображение.")
        
        img_pil = Image.open(self.image_path).convert("L")
        img_pil = img_pil.convert("RGB")
        img_tensor = self.transforms(img_pil) if self.transforms else T.ToTensor()(img_pil)
        img_name = os.path.basename(self.image_path)
        
        return img_tensor, img_pil, img_name

# Трансформ для данных (Faster R-CNN)
def get_test_transform():
    return T.Compose([
        T.ToTensor(),
    ])

# ---------------------------------------------------------------------------------
# Модель TransformerOCR – та же структура и для текста, и для математики
# ---------------------------------------------------------------------------------
class PositionalEncoding(torch.nn.Module):
    def __init__(self, d_model, dropout=0.1, max_len=5000):
        super(PositionalEncoding, self).__init__()
        self.dropout = torch.nn.Dropout(p=dropout)
        self.scale = torch.nn.Parameter(torch.ones(1))

        pe = torch.zeros(max_len, d_model)
        position = torch.arange(0, max_len, dtype=torch.float).unsqueeze(1)
        div_term = torch.exp(torch.arange(0, d_model, 2).float() 
                             * (-math.log(10000.0) / d_model))
        pe[:, 0::2] = torch.sin(position * div_term)
        pe[:, 1::2] = torch.cos(position * div_term)
        pe = pe.unsqueeze(0).transpose(0, 1)  # [max_len, 1, d_model]
        self.register_buffer('pe', pe)

    def forward(self, x):
        x = x + self.scale * self.pe[:x.size(0), :]
        return self.dropout(x) 

class TransformerModelOCR(nn.Module):
    def __init__(self, outtoken, hidden, enc_layers=1, dec_layers=1, nhead=1, dropout=0.1):
        super(TransformerModelOCR, self).__init__()
        self.conv0 = Conv2d(1, 64, kernel_size=(3, 3), stride=(1, 1), padding=(1, 1))
        self.conv1 = Conv2d(64, 128, kernel_size=(3, 3), stride=(1, 1), padding=(1, 1))
        self.conv2 = Conv2d(128, 256, kernel_size=(3, 3), stride=(2, 1), padding=(1, 1))
        self.conv3 = Conv2d(256, 256, kernel_size=(3, 3), stride=(1, 1), padding=(1, 1))
        self.conv4 = Conv2d(256, 512, kernel_size=(3, 3), stride=(2, 1), padding=(1, 1))
        self.conv5 = Conv2d(512, 512, kernel_size=(3, 3), stride=(1, 1), padding=(1, 1))
        self.conv6 = Conv2d(512, 512, kernel_size=(2, 1), stride=(1, 1))

        self.pool1 = MaxPool2d(kernel_size=2, stride=2)
        self.pool3 = MaxPool2d(kernel_size=2, stride=2)
        self.pool5 = MaxPool2d(kernel_size=(2, 2), stride=(2, 1), padding=(0, 1))

        self.bn0 = BatchNorm2d(64)
        self.bn1 = BatchNorm2d(128)
        self.bn2 = BatchNorm2d(256)
        self.bn3 = BatchNorm2d(256)
        self.bn4 = BatchNorm2d(512)
        self.bn5 = BatchNorm2d(512)
        self.bn6 = BatchNorm2d(512)

        self.activ = LeakyReLU()

        self.pos_encoder = PositionalEncoding(hidden, dropout)
        self.decoder = nn.Embedding(outtoken, hidden)
        self.pos_decoder = PositionalEncoding(hidden, dropout)
        self.transformer = nn.Transformer(d_model=hidden, nhead=nhead, 
                                          num_encoder_layers=enc_layers,
                                          num_decoder_layers=dec_layers, 
                                          dim_feedforward=hidden*4, 
                                          dropout=dropout)
        self.fc_out = nn.Linear(hidden, outtoken)
        self.src_mask = None
        self.trg_mask = None
        self.memory_mask = None

    def generate_square_subsequent_mask(self, sz):
        mask = torch.triu(torch.ones(sz, sz, device=DEVICE), 1)
        mask = mask.masked_fill(mask == 1, float('-inf'))
        return mask

    def make_len_mask(self, inp):
        return (inp == 0).transpose(0, 1)
    
    def _get_features(self, src):
        x = self.activ(self.bn0(self.conv0(src)))
        x = self.pool1(self.activ(self.bn1(self.conv1(x))))
        x = self.activ(self.bn2(self.conv2(x)))
        x = self.pool3(self.activ(self.bn3(self.conv3(x))))
        x = self.activ(self.bn4(self.conv4(x)))
        x = self.pool5(self.activ(self.bn5(self.conv5(x))))
        x = self.activ(self.bn6(self.conv6(x)))
        # Преобразуем в (sequence_len, batch, channels)
        x = x.permute(0, 3, 1, 2).flatten(2).permute(1, 0, 2)
        return x

    def predict(self, batch, char2idx):
        results = []
        for item in batch:  # итерация по батчу
            x = self._get_features(item.unsqueeze(0))
            memory = self.transformer.encoder(self.pos_encoder(x))
            out_indexes = [char2idx['SOS']]
            for _ in range(300):  # некий лимит
                trg_tensor = torch.LongTensor(out_indexes).unsqueeze(1).to(DEVICE)
                output = self.fc_out(
                    self.transformer.decoder(
                        self.pos_decoder(self.decoder(trg_tensor)), 
                        memory
                    )
                )
                out_token = output.argmax(2)[-1].item()
                out_indexes.append(out_token)
                if 'EOS' in char2idx and out_token == char2idx['EOS']:
                    break
            results.append(out_indexes)
        return results

    def forward(self, src, trg):
        if self.trg_mask is None or self.trg_mask.size(0) != len(trg):
            self.trg_mask = self.generate_square_subsequent_mask(len(trg)).to(trg.device)
        x = self._get_features(src)
        src_pad_mask = self.make_len_mask(x[:, :, 0])
        src = self.pos_encoder(x)
        trg_pad_mask = self.make_len_mask(trg)
        trg = self.decoder(trg)
        trg = self.pos_decoder(trg)
        output = self.transformer(
            src, trg, src_mask=self.src_mask, tgt_mask=self.trg_mask,
            memory_mask=self.memory_mask,
            src_key_padding_mask=src_pad_mask, 
            tgt_key_padding_mask=trg_pad_mask,
            memory_key_padding_mask=src_pad_mask
        )
        output = self.fc_out(output)
        return output

# ---------------------------------------------------------------------------------
# Универсальная функция загрузки OCR (текст или математика)
# ---------------------------------------------------------------------------------
def load_ocr_model(weights_path, alphabet):
    """
    Создаёт модель TransformerModelOCR с размером словаря = len(alphabet),
    загружает веса из weights_path и возвращает (model, char2idx, idx2char).
    """
    hidden = 512
    enc_layers = 2
    dec_layers = 2
    n_heads = 4

    char2idx = {c: i for i, c in enumerate(alphabet)}
    idx2char = {i: c for i, c in enumerate(alphabet)}

    model_ocr = TransformerModelOCR(
        outtoken=len(alphabet),
        hidden=hidden,
        enc_layers=enc_layers,
        dec_layers=dec_layers,
        nhead=n_heads,
        dropout=0.2
    ).to(DEVICE)

    checkpoint = torch.load(weights_path, map_location=DEVICE, weights_only=True)
    model_ocr.load_state_dict(checkpoint)
    model_ocr.eval()
    print(f"Модель {weights_path} загружена. Размер алфавита = {len(alphabet)}.")

    return model_ocr, char2idx, idx2char

# Загружаем отдельно модели и словари
model_ocr, char2idx_text, idx2char_text = load_ocr_model(WEIGHTS_PATH_TEXT, ALPHABET_TEXT)
model_math, char2idx_math, idx2char_math = load_ocr_model(WEIGHTS_PATH_MATH, ALPHABET_MATH)

# ---------------------------------------------------------------------------------
# Функции сортировки/группировки сегментов
# ---------------------------------------------------------------------------------
def sort_segments(segments, y_threshold=30, x_threshold=20):
    segments_sorted = sorted(segments, key=lambda seg: seg['coords'][1])
    grouped_segments = []
    current_group = []

    for seg in segments_sorted:
        if not current_group:
            current_group.append(seg)
        else:
            last_seg = current_group[-1]
            if abs(seg['coords'][1] - last_seg['coords'][1]) <= y_threshold:
                current_group.append(seg)
            else:
                grouped_segments.append(current_group)
                current_group = [seg]
    if current_group:
        grouped_segments.append(current_group)

    final_segments = []
    for group in grouped_segments:
        group_sorted = sorted(group, key=lambda seg: seg['coords'][0])
        final_segments.extend(group_sorted)
    return final_segments

def group_all_segments_into_lines(segments, y_threshold=10):
    if not segments:
        return []
    grouped = []
    current_line = [segments[0]]
    current_ymin = segments[0]['coords'][1]
    current_ymax = segments[0]['coords'][3]

    for seg in segments[1:]:
        ymin, ymax = seg['coords'][1], seg['coords'][3]
        if abs(ymin - current_ymin) <= y_threshold:
            current_line.append(seg)
            current_ymin = min(current_ymin, ymin)
            current_ymax = max(current_ymax, ymax)
        else:
            grouped.append(current_line)
            current_line = [seg]
            current_ymin = ymin
            current_ymax = ymax

    grouped.append(current_line)
    return grouped

def sort_segments_within_line(line):
    return sorted(line, key=lambda seg: seg['coords'][0])

# ---------------------------------------------------------------------------------
# Распознавание сегмента (Text / Math)
# ---------------------------------------------------------------------------------
def indices_to_text(indexes, idx2char):
    """
    Преобразует список индексов в строку, убирая SOS/EOS.
    """
    result = []
    for idx in indexes:
        if idx2char[idx] in ("SOS", "EOS", "PAD"):
            continue
        result.append(idx2char[idx])
    return "".join(result)

def recognize_segment(segment_img, label):
    """
    Распознаёт сегмент в зависимости от метки (Text/Math).
    Возвращает строку с результатом.
    """
    if segment_img is None:
        return ""

    # Преобразуем в нужный формат
    img_pil = segment_img.convert('L')
    img_transformed = transforms.Compose([
        transforms.Resize((64, 256)),
        transforms.ToTensor(),
        transforms.Normalize((0.5,), (0.5,))
    ])(img_pil).unsqueeze(0).to(DEVICE)

    # Выбираем нужную модель и словари
    if label == "Text":
        model = model_ocr
        c2i = char2idx_text
        i2c = idx2char_text
    elif label == "Math":
        model = model_math
        c2i = char2idx_math
        i2c = idx2char_math
    else:
        return ""  # например, для Image

    with torch.no_grad():
        out_indexes = model.predict(img_transformed, c2i)  # список списков индексов
    recognized_text = indices_to_text(out_indexes[0], i2c)
    return recognized_text

# ---------------------------------------------------------------------------------
# Основная функция для детекции и распознавания
# ---------------------------------------------------------------------------------
def predict_and_recognize_text(test_dataset, 
                               model_faster,
                               device,
                               num_images=5, 
                               score_threshold=0.3):
    model_faster.eval()
    dataset_size = len(test_dataset)
    if dataset_size == 0:
        print("Тестовый датасет пуст.")
        return
    
    num_images = min(num_images, dataset_size)
    random_indices = random.sample(range(dataset_size), num_images)

    document = Document()

    for idx in random_indices:
        img_tensor, img_pil_original, img_name = test_dataset[idx]
        # Сразу работаем с тензором, PIL нам нужен только для вырезания сегментов
        # (Если вырезание делаете через PIL, как в коде ниже)

        img_input = img_tensor.to(device).unsqueeze(0)
        with torch.no_grad():
            prediction = model_faster(img_input)
        
        boxes = prediction[0]['boxes'].cpu().numpy()
        labels = prediction[0]['labels'].cpu().numpy()
        scores = prediction[0]['scores'].cpu().numpy()

        # Фильтруем по порогу уверенности
        keep = scores >= score_threshold
        boxes = boxes[keep]
        labels = labels[keep]
        scores = scores[keep]

        if len(boxes) == 0:
            print(f"Нет объектов выше порога {score_threshold} на {img_name}.")
            continue

        # Список сегментов
        segments = []
        for box, label, score_val in zip(boxes, labels, scores):
            xmin, ymin, xmax, ymax = map(int, box)
            seg_img = None
            if CLASS_NAMES.get(label) in ["Text", "Math", "Image"]:
                seg_img = img_pil_original.crop((xmin, ymin, xmax, ymax))
            
            seg_info = {
                'image': seg_img,
                'coords': (xmin, ymin, xmax, ymax),
                'label': CLASS_NAMES.get(label, "Unknown"),
                'score': float(score_val),
            }
            segments.append(seg_info)

        # Сортируем и группируем сегменты
        segments_sorted = sort_segments(segments)
        grouped_segments = group_all_segments_into_lines(segments_sorted, y_threshold=100)

        original_width, original_height = img_pil_original.size

        # Проходим по сгруппированным «строкам»
        for line in grouped_segments:
            sorted_line = sort_segments_within_line(line)
            paragraph = document.add_paragraph()

            for elem in sorted_line:
                if elem['label'] in ["Text", "Math"]:
                    recognized_text = recognize_segment(elem['image'], elem['label'])
                    paragraph.add_run(recognized_text + " ")
                elif elem['label'] == "Image" and elem['image'] is not None:
                    xmin, ymin, xmax, ymax = elem['coords']
                    seg_w = xmax - xmin
                    seg_h = ymax - ymin
                    max_width_inches = 6
                    relative_width = seg_w / original_width
                    desired_w_inches = relative_width * max_width_inches
                    desired_w_inches = min(desired_w_inches, max_width_inches)

                    try:
                        seg_img_resized = elem['image'].resize(
                            (int(seg_w * (desired_w_inches / max_width_inches)),
                             int(seg_h * (desired_w_inches / max_width_inches))),
                            Image.Resampling.LANCZOS
                        )
                    except AttributeError:
                        seg_img_resized = elem['image'].resize(
                            (int(seg_w * (desired_w_inches / max_width_inches)),
                             int(seg_h * (desired_w_inches / max_width_inches))),
                            Image.LANCZOS
                        )
                    
                    buffered_seg = BytesIO()
                    seg_img_resized.save(buffered_seg, format="PNG")
                    buffered_seg.seek(0)

                    run = paragraph.add_run()
                    run.add_picture(buffered_seg, width=Inches(desired_w_inches))
                    run.add_text(" ")

    # Сохраняем результат в единый документ 'OCR.docx'
    doc_path = os.path.join(OUTPUT_FOLDER, "OCR.docx")
    document.save(doc_path)

    return document

# ---------------------------------------------------------------------------------
# FastAPI
# ---------------------------------------------------------------------------------
app = FastAPI()

@app.get("/")
async def root():
    return {"message": "FastAPI OCR and Object Detection"}

@app.post("/upload_image/")
async def upload_image(file: UploadFile = File(...)):
    try:
        file_path = os.path.join(UPLOAD_FOLDER, file.filename)
        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)

        test_dataset = TestImageDataset(file_path, transforms=get_test_transform())
        if len(test_dataset) == 0:
            raise HTTPException(status_code=400, detail="Пустой датасет")

        document = predict_and_recognize_text(
            test_dataset=test_dataset,
            model_faster=model_faster,
            device=DEVICE,
            num_images=1,            # т.к. dataset всё равно 1
            score_threshold=0.01
        )
        if not document:
            raise HTTPException(status_code=500, detail="Ошибка при распознавании")

        # Можно вернуть пути к файлам
        
        return {"ocr_result": "OCR.docx"}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/download_file/{file_name}")
async def download_file(file_name: str):
    file_path = Path(OUTPUT_FOLDER) / file_name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(file_path, media_type="application/octet-stream", 
                       headers={"Content-Disposition": f"attachment; filename={file_name}"})
